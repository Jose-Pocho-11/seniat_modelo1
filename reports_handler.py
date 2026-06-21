import sqlite3
import os
from datetime import datetime
from init_db import DB_PATH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.units import inch
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_periodo(periodo_str):
    if not periodo_str: return None
    s = str(periodo_str).strip()
    if ' ' in s: s = s.split(' ')[0]
    parts = s.replace('/', '-').split('-')
    if len(parts) == 2:
        if len(parts[1]) == 4:
            y = 2000 + int(parts[1][:2])
            m = int(parts[1][2:4])
            return str(y), str(m)
        p0, p1 = int(parts[0]), int(parts[1])
        if p0 > 1000: return str(p0), str(p1)
        return str(p1), str(p0)
    if len(parts) >= 3:
        p0, p1, p2 = int(parts[0]), int(parts[1]), int(parts[2])
        if p0 > 31: return str(p0), str(p1)
        return str(p2), str(p1)
    return None, None

def check_compliance_py(rif_str, periodo_str, fecha_pago_str, impuesto_str, cal_cache, cal_dpp_cache):
    if not rif_str or not periodo_str or not fecha_pago_str: return None
    last_num = str(rif_str).strip()[-1:]
    if not last_num.isdigit(): return None
    terminal = int(last_num)
    
    y_per, m_per = parse_periodo(periodo_str)
    if not y_per or not m_per: return None
    
    try:
        f_str = str(fecha_pago_str).strip().split(' ')[0]
        if '-' in f_str:
            dPago = datetime.strptime(f_str, "%Y-%m-%d")
        else:
            dPago = datetime.strptime(f_str, "%d/%m/%Y")
    except:
        return None

    dueMonth = int(m_per) - 1
    dueYear = int(y_per)
    
    is_dpp = 'dpp' in str(impuesto_str).lower() or 'anticipo' in str(impuesto_str).lower()
    months = ['ENE', 'FEB', 'MAR', 'ABR', 'MAY', 'JUN', 'JUL', 'AGO', 'SEP', 'OCT', 'NOV', 'DIC']
    dueDay = None
    
    if not is_dpp:
        quincena = 2 if dPago.day <= 15 else 1
        if quincena == 2:
            dueMonth += 1
            if dueMonth > 11:
                dueMonth = 0
                dueYear += 1
        monthStrLookup = months[dueMonth]
        try:
            dueDay = cal_cache[dueYear][quincena][monthStrLookup][terminal]
        except KeyError:
            return None
    else:
        dueMonth += 1
        if dueMonth > 11:
            dueMonth = 0
            dueYear += 1
        monthStrLookup = months[dueMonth]
        try:
            dueDay = cal_dpp_cache[dueYear][monthStrLookup][terminal]
        except KeyError:
            return None
            
    if not dueDay: return None
    dDue = datetime(dueYear, dueMonth + 1, dueDay, 23, 59, 59)
    dPago = dPago.replace(hour=0, minute=0, second=0)
    
    return dPago <= dDue

def get_report_data(user_id, is_admin, year=None, month=None):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM calendarios')
    cals = cursor.fetchall()
    calendar_cache = {}
    for r in cals:
        y, q, m, t, d = r['anio'], r['quincena'], r['mes'], r['terminal_rif'], r['dia_vencimiento']
        if y not in calendar_cache: calendar_cache[y] = {1: {}, 2: {}}
        if m not in calendar_cache[y][q]: calendar_cache[y][q][m] = {}
        calendar_cache[y][q][m][t] = d

    cursor.execute('SELECT * FROM calendarios_dpp')
    cals_dpp = cursor.fetchall()
    calendar_dpp_cache = {}
    for r in cals_dpp:
        y, m, t, d = r['anio'], r['mes'], r['terminal_rif'], r['dia_vencimiento']
        if y not in calendar_dpp_cache: calendar_dpp_cache[y] = {}
        if m not in calendar_dpp_cache[y]: calendar_dpp_cache[y][m] = {}
        calendar_dpp_cache[y][m][t] = d
    
    if is_admin:
        query = '''
            SELECT r.id, r.monto, r.rif, r.razon_social, r.periodo, r.fecha_recaudacion,
                   d.nombre_dependencia as dependencia,
                   b.nombre_banco as banco,
                   ti.nombre_impuesto as impuesto
            FROM recaudaciones r
            JOIN dependencias d ON r.dependencia_id = d.id
            JOIN bancos b ON r.banco_id = b.id
            JOIN tipos_impuesto ti ON r.impuesto_id = ti.id
        '''
        params = []
    else:
        query = '''
            SELECT r.id, r.monto, r.rif, r.razon_social, r.periodo, r.fecha_recaudacion,
                   d.nombre_dependencia as dependencia,
                   b.nombre_banco as banco,
                   ti.nombre_impuesto as impuesto
            FROM recaudaciones r
            JOIN dependencias d ON r.dependencia_id = d.id
            JOIN bancos b ON r.banco_id = b.id
            JOIN tipos_impuesto ti ON r.impuesto_id = ti.id
            JOIN historial_lotes h ON r.id = h.id
            WHERE h.usuario_id = ?
        '''
        params = [user_id]
        
    cursor.execute(query, params)
    rows = cursor.fetchall()
    conn.close()
    
    total = 0.0
    cantidad_pagos = 0
    dep_dict = {}
    ban_dict = {}
    imp_dict = {}
    term_dict = {}
    cont_dict = {}
    comp_dpp_dict = {str(i): {'a_tiempo': 0, 'retrasado': 0} for i in range(10)}
    comp_otros_dict = {str(i): {'a_tiempo': 0, 'retrasado': 0} for i in range(10)}
    
    filter_year = str(int(year)) if year else None
    filter_month = str(int(month)) if month else None
    
    for row in rows:
        r_year, r_month = parse_periodo(row['periodo'])
        
        if filter_year and r_year != filter_year: continue
        if filter_month and r_month != filter_month: continue
            
        monto = float(row['monto'])
        total += monto
        cantidad_pagos += 1
        
        dep = row['dependencia']
        dep_dict[dep] = dep_dict.get(dep, 0) + monto
        
        ban = row['banco']
        ban_dict[ban] = ban_dict.get(ban, 0) + monto
        
        imp = row['impuesto']
        imp_dict[imp] = imp_dict.get(imp, 0) + monto
        
        rif = str(row['rif'])
        terminal = rif[-1] if rif else ''
        term_dict[terminal] = term_dict.get(terminal, 0) + monto
        
        razon = str(row['razon_social'])
        c_key = (rif, razon)
        cont_dict[c_key] = cont_dict.get(c_key, 0) + monto
        
        # Compliance Check
        if terminal.isdigit():
            is_dpp = 'dpp' in str(imp).lower() or 'anticipo' in str(imp).lower()
            is_on_time = check_compliance_py(rif, row['periodo'], row['fecha_recaudacion'], imp, calendar_cache, calendar_dpp_cache)
            if is_on_time is True:
                if is_dpp:
                    comp_dpp_dict[terminal]['a_tiempo'] += 1
                else:
                    comp_otros_dict[terminal]['a_tiempo'] += 1
            elif is_on_time is False:
                if is_dpp:
                    comp_dpp_dict[terminal]['retrasado'] += 1
                else:
                    comp_otros_dict[terminal]['retrasado'] += 1

    dependencias = [{"nombre": k, "monto": v} for k, v in sorted(dep_dict.items(), key=lambda x: x[1], reverse=True)]
    bancos = [{"nombre": k, "monto": v} for k, v in sorted(ban_dict.items(), key=lambda x: x[1], reverse=True)]
    impuestos = [{"nombre": k, "monto": v} for k, v in sorted(imp_dict.items(), key=lambda x: x[1], reverse=True)]
    terminales = [{"terminal": k, "monto": v} for k, v in sorted(term_dict.items(), key=lambda x: x[1], reverse=True)]
    top_contribuyentes = [{"rif": k[0], "razon_social": k[1], "monto": v} for k, v in sorted(cont_dict.items(), key=lambda x: x[1], reverse=True)]
    
    def build_comp_list(d):
        lst = []
        for term in range(10):
            t_str = str(term)
            atiempo = d[t_str]['a_tiempo']
            retraso = d[t_str]['retrasado']
            total_eval = atiempo + retraso
            if total_eval > 0:
                perc_atiempo = (atiempo / total_eval) * 100
                perc_retraso = (retraso / total_eval) * 100
            else:
                perc_atiempo = 0
                perc_retraso = 0
            lst.append({
                "terminal": t_str,
                "evaluadas": total_eval,
                "a_tiempo_pct": perc_atiempo,
                "retrasado_pct": perc_retraso
            })
        return lst

    compliance_dpp_list = build_comp_list(comp_dpp_dict)
    compliance_otros_list = build_comp_list(comp_otros_dict)
    
    return {
        "total": total,
        "cantidad_de_pagos": cantidad_pagos,
        "dependencias": dependencias,
        "bancos": bancos,
        "impuestos": impuestos,
        "terminales": terminales,
        "top_contribuyentes": top_contribuyentes,
        "compliance_dpp": compliance_dpp_list,
        "compliance_otros": compliance_otros_list
    }

def format_bs(amount):
    return f"Bs. {amount:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def generate_pdf_report(user_id, is_admin, output_path, year=None, month=None):
    data = get_report_data(user_id, is_admin, year, month)
    
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(name='TitleStyle', parent=styles['Heading1'], alignment=1, fontSize=16, spaceAfter=20)
    subtitle_style = ParagraphStyle(name='SubtitleStyle', parent=styles['Heading2'], fontSize=12, spaceAfter=10, textColor=colors.HexColor('#1e40af'))
    normal_style = styles['Normal']
    
    logo_path = os.path.join(os.path.dirname(__file__), 'static', 'img', 'logo_seniat.png')
    if os.path.exists(logo_path):
        img = Image(logo_path, width=2*inch, height=0.6*inch)
        elements.append(img)
        elements.append(Spacer(1, 10))
    
    elements.append(Paragraph("<b>Reporte Oficial de Recaudación</b>", title_style))
    elements.append(Paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", normal_style))
    elements.append(Spacer(1, 20))
    
    elements.append(Paragraph(f"<b>Monto Total Recaudado:</b> {format_bs(data['total'])}", styles['Heading3']))
    elements.append(Paragraph(f"<b>Cantidad de Pagos:</b> {data['cantidad_de_pagos']} pagos", normal_style))
    elements.append(Spacer(1, 15))
    
    def add_table_section(title, headers, rows):
        if not rows:
            return
        elements.append(Paragraph(f"<b>{title}</b>", subtitle_style))
            
        table_data = [headers]
        for r in rows:
            name_text = str(r['nombre'] if 'nombre' in r else r['terminal'])
            name_paragraph = Paragraph(name_text, normal_style)
            table_data.append([name_paragraph, format_bs(r['monto'])])
            
        t = Table(table_data, colWidths=[4*inch, 2.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e2e8f0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 20))
        if rows:
            top_item = rows[0]
            top_name = str(top_item['nombre'] if 'nombre' in top_item else top_item['terminal'])
            top_monto = format_bs(top_item['monto'])
            prose = f"Como se observa en el desglose, la mayor parte de la recaudación en esta categoría proviene de <b>{top_name}</b>, aportando un total de {top_monto}. Esto refleja la principal fuente de ingresos en el período consultado."
            elements.append(Paragraph(prose, normal_style))
            elements.append(Spacer(1, 20))

    add_table_section("1. Desglose por Dependencia", ["Dependencia", "Monto"], data['dependencias'])
    add_table_section("2. Desglose por Tipos de Impuesto", ["Impuesto", "Monto"], data['impuestos'])
    add_table_section("3. Desglose por Banco", ["Banco", "Monto"], data['bancos'])
    add_table_section("4. Recaudación por Terminal de RIF", ["Terminal de RIF", "Monto"], data['terminales'])
    
    def add_compliance_table(title, comp_rows):
        valid_comp = [r for r in comp_rows if r['evaluadas'] > 0]
        if not valid_comp:
            return
            
        elements.append(Paragraph(f"<b>{title}</b>", subtitle_style))
        table_data = [["Terminal", "Evaluadas", "A Tiempo (%)", "Retrasado (%)"]]
        for r in valid_comp:
            t_val = r['terminal']
            e = r['evaluadas']
            at = f"{r['a_tiempo_pct']:.1f}%"
            rt = f"{r['retrasado_pct']:.1f}%"
            table_data.append([t_val, str(e), at, rt])
            
        t = Table(table_data, colWidths=[1.5*inch, 1.5*inch, 1.75*inch, 1.75*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e2e8f0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0f172a')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#cbd5e1')),
            ('TEXTCOLOR', (2, 1), (2, -1), colors.HexColor('#15803d')),
            ('TEXTCOLOR', (3, 1), (3, -1), colors.HexColor('#b91c1c')),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 20))

    # 5. Comportamiento de Pago
    add_compliance_table("5. Comportamiento de Pago por Terminal de RIF (Anticipos DPP)", data.get('compliance_dpp', []))
    add_compliance_table("6. Comportamiento de Pago por Terminal de RIF (Retenciones / Otros)", data.get('compliance_otros', []))
        
    doc.build(elements)
